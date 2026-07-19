import os
import json
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

RESULTS_DIR = 'results'
DOC_PATH = os.path.join(RESULTS_DIR, 'PQC_Benchmarks_Report.docx')

def load_json(filename):
    path = os.path.join(RESULTS_DIR, filename)
    if os.path.exists(path):
        with open(path, 'r') as f:
            return json.load(f)
    return None

def add_table_from_json(doc, data, columns_to_show):
    if not data or 'results' not in data:
        p = doc.add_paragraph()
        p.add_run("No data available.").italic = True
        return

    results = data['results']
    if not results:
        p = doc.add_paragraph()
        p.add_run("No data available.").italic = True
        return

    # Add table
    table = doc.add_table(rows=1, cols=len(columns_to_show))
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(columns_to_show.values()):
        hdr_cells[i].text = col_name
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    # Data rows
    for row in results:
        row_cells = table.add_row().cells
        for i, key in enumerate(columns_to_show.keys()):
            val = row.get(key, 'N/A')
            if isinstance(val, float):
                val = f"{val:.2f}"
            row_cells[i].text = str(val)

def generate_word_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('XenevaOS Post-Quantum Cryptography Benchmarks', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document contains the finalized data extractions from the cryptographic simulation experiments for XenevaOS. All benchmarks were run natively in C++ using OpenSSL 3.5.5.')
    
    # System Info
    sys_info = {}
    data1 = load_json('ecdh_vs_mlkem_results.json')
    if data1 and 'system' in data1:
        sys_info = data1['system']
        
    doc.add_heading('System Configuration', level=1)
    for k, v in sys_info.items():
        doc.add_paragraph(f"{k.capitalize()}: {v}", style='List Bullet')
        
    # Standard columns
    std_cols = {
        'algorithm': 'Algorithm',
        'operation': 'Operation',
        'average': 'Mean Latency (µs)',
        'ops_per_sec': 'Ops/sec',
        'relative_slowdown': 'Slowdown (x)'
    }
        
    # Experiment 1
    doc.add_heading('1. ECDH P-256 vs ML-KEM-768 (Key Establishment)', level=1)
    doc.add_paragraph('Compares the latency of generating keys and establishing shared secrets.')
    add_table_from_json(doc, load_json('ecdh_vs_mlkem_results.json'), std_cols)
    if os.path.exists(os.path.join(RESULTS_DIR, 'plots', 'exp1_ecdh_vs_mlkem.png')):
        doc.add_picture(os.path.join(RESULTS_DIR, 'plots', 'exp1_ecdh_vs_mlkem.png'), width=Inches(5.5))
        
    # Experiment 2
    doc.add_page_break()
    doc.add_heading('2. RSA-2048 vs ML-DSA-65 (Digital Signatures)', level=1)
    doc.add_paragraph('Compares standard key/signature generation and verification. Includes a 256-byte payload.')
    add_table_from_json(doc, load_json('rsa_vs_mldsa_results.json'), std_cols)
    if os.path.exists(os.path.join(RESULTS_DIR, 'plots', 'exp2_rsa_vs_mldsa.png')):
        doc.add_picture(os.path.join(RESULTS_DIR, 'plots', 'exp2_rsa_vs_mldsa.png'), width=Inches(5.5))
        
    # Experiment 3
    doc.add_page_break()
    doc.add_heading('3. Secure Boot Simulation (Kernel Integrity)', level=1)
    doc.add_paragraph('Simulates the impact of signature verification on OS boot times using a synthetic 4MB kernel image payload.')
    add_table_from_json(doc, load_json('secure_boot_results.json'), std_cols)
    if os.path.exists(os.path.join(RESULTS_DIR, 'plots', 'exp3_secure_boot.png')):
        doc.add_picture(os.path.join(RESULTS_DIR, 'plots', 'exp3_secure_boot.png'), width=Inches(5.5))
        
    # Experiment 4
    doc.add_page_break()
    doc.add_heading('4. Package Verification (Software Supply Chain)', level=1)
    doc.add_paragraph('Simulates verifying batches of 1MB system packages (100, 500, and 1000 packages per batch).')
    add_table_from_json(doc, load_json('package_verify_results.json'), std_cols)
    if os.path.exists(os.path.join(RESULTS_DIR, 'plots', 'exp4_package_verify.png')):
        doc.add_picture(os.path.join(RESULTS_DIR, 'plots', 'exp4_package_verify.png'), width=Inches(5.5))
        
    # Experiment 5
    doc.add_page_break()
    doc.add_heading('5. Secure IPC Simulation (Process Isolation)', level=1)
    doc.add_paragraph('Simulates a full inter-process communication handshake: KeyGen -> KeyDerivation -> Encrypt/Decrypt (AES-256-GCM) on a 4KB payload.')
    add_table_from_json(doc, load_json('secure_ipc_results.json'), std_cols)
    if os.path.exists(os.path.join(RESULTS_DIR, 'plots', 'exp5_secure_ipc.png')):
        doc.add_picture(os.path.join(RESULTS_DIR, 'plots', 'exp5_secure_ipc.png'), width=Inches(5.5))
        
    doc.save(DOC_PATH)
    print(f"Successfully generated Word Document at {DOC_PATH}")

if __name__ == '__main__':
    generate_word_doc()
